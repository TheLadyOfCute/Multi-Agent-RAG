"""
Document Loader - Load and extract text from various file formats.

Supports: PDF, DOCX, TXT, MD
"""

from typing import Dict, Any
from pathlib import Path
import hashlib
from datetime import datetime

from src.utils.logger import setup_logger
from src.utils.exceptions import AgenticRAGException


class DocumentLoadError(AgenticRAGException):
    """Error during document loading."""
    pass


class Document:
    def __init__(self, text: str, metadata: Dict[str, Any] = None):
        self.text = text
        self.metadata = metadata or {}
        self.doc_id = self._generate_doc_id()

        if "loaded_at" not in self.metadata:
            self.metadata["loaded_at"] = datetime.now().isoformat()

    def _generate_doc_id(self) -> str:
        """Generate unique document ID from text hash."""
        return hashlib.md5(self.text.encode()).hexdigest()

    def __len__(self) -> int:
        """Get document length."""
        return len(self.text)

    def __repr__(self) -> str:
        """String representation."""
        return f"Document(doc_id={self.doc_id[:8]}..., length={len(self.text)})"


class DocumentLoader:
    def __init__(self):
        """Initialize document loader."""
        self.logger = setup_logger("document_loader")
        self.supported_formats = [".pdf", ".docx", ".txt", ".md"]

    def load(self, file_path: str) -> Document:
        path = Path(file_path)

        if not path.exists():
            raise DocumentLoadError(
                message=f"File not found: {file_path}",
                details={"path": file_path},
            )

        file_ext = path.suffix.lower()

        if file_ext not in self.supported_formats:
            raise DocumentLoadError(
                message=f"Unsupported file format: {file_ext}",
                details={"format": file_ext, "supported": self.supported_formats},
            )

        self.logger.info(f"Loading document: {path.name}")

        try:
            if file_ext == ".pdf":
                text = self._load_pdf(path)
            elif file_ext == ".docx":
                text = self._load_docx(path)
            else:  # .txt, .md
                text = self._load_text(path)

            metadata = self._extract_metadata(path, file_ext)
            doc = Document(text=text, metadata=metadata)

            self.logger.info(
                f"Loaded document: {path.name} "
                f"({len(doc.text)} chars, doc_id={doc.doc_id[:8]}...)"
            )
            return doc

        except Exception as e:
            self.logger.error(f"Failed to load {path.name}: {str(e)}")
            raise DocumentLoadError(
                message=f"Failed to load document: {str(e)}",
                details={"path": file_path, "error": str(e)},
            ) from e

    def count_pages(self, file_path: str) -> int:
        path = Path(file_path)

        if not path.exists():
            self.logger.warning(f"File not found: {file_path}")
            return 1

        file_ext = path.suffix.lower()

        try:
            if file_ext == ".pdf":
                from pypdf import PdfReader

                reader = PdfReader(path)
                return len(reader.pages)

            if file_ext == ".docx":
                from docx import Document as DocxDocument

                doc = DocxDocument(path)
                paragraph_count = len(doc.paragraphs)
                return max(1, paragraph_count // 30)

            # .txt, .md
            with open(path, "r", encoding="utf-8") as f:
                line_count = len(f.readlines())
            return max(1, line_count // 40)

        except Exception as e:
            self.logger.warning(f"Could not count pages for {path.name}: {e}")
            return 1

    def _load_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise DocumentLoadError(
                message="pypdf not installed. Run: pip install pypdf",
                details={"required_package": "pypdf"},
            )

        text_parts = []
        reader = PdfReader(path)

        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)

        if not text_parts:
            raise DocumentLoadError(
                message="No text content found in PDF",
                details={"path": str(path)},
            )

        return "\n\n".join(text_parts)

    def _load_docx(self, path: Path) -> str:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise DocumentLoadError(
                message="python-docx not installed. Run: pip install python-docx",
                details={"required_package": "python-docx"},
            )

        doc = DocxDocument(path)
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]

        if not text_parts:
            raise DocumentLoadError(
                message="No text content found in DOCX",
                details={"path": str(path)},
            )

        return "\n\n".join(text_parts)

    def _load_text(self, path: Path) -> str:
        try:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()

            if not text.strip():
                raise DocumentLoadError(
                    message="File is empty",
                    details={"path": str(path)},
                )

            return text

        except UnicodeDecodeError:
            try:
                with open(path, "r", encoding="latin-1") as f:
                    return f.read()
            except Exception as e:
                raise DocumentLoadError(
                    message=f"Failed to decode text file: {str(e)}",
                    details={"path": str(path)},
                )

    def _extract_metadata(self, path: Path, file_ext: str) -> Dict[str, Any]:
        stat = path.stat()

        return {
            "filename": path.name,
            "file_path": str(path.absolute()),
            "file_type": file_ext,
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "source": "file_upload",
        }

